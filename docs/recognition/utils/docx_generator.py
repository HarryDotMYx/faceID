import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.conf import settings
import logging

logger = logging.getLogger(__name__)

class DocxGenerator:
    def __init__(self):
        self.document = Document()
        self._setup_document()
        
    def _setup_document(self):
        """Setup document styles"""
        style = self.document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        
    def add_title(self, title):
        """Add document title"""
        heading = self.document.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_paragraph()
        
    def add_section(self, title, content):
        """Add a new section with title and content"""
        self.document.add_heading(title, level=1)
        self.document.add_paragraph(content)
        self.document.add_paragraph()
        
    def add_image(self, image_path, width=6):
        """Add an image to the document"""
        try:
            self.document.add_picture(image_path, width=Inches(width))
            self.document.add_paragraph()
        except Exception as e:
            logger.error(f"Error adding image to document: {str(e)}")
            
    def add_table(self, headers, rows):
        """Add a table to the document"""
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            
        # Add data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                
        self.document.add_paragraph()
        
    def save(self, filename):
        """Save document to file"""
        try:
            filepath = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            self.document.save(filepath)
            return filepath
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return None